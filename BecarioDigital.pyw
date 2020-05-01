from docx import Document
import os
import win32api
import pyforms
import time
from datetime import datetime,timedelta,date
from pyforms import settings as formSettings
from pyforms.basewidget import BaseWidget
from pyforms.controls   import ControlText
from pyforms.controls   import ControlButton
from pyforms.controls   import ControlCheckBox
from pyforms.controls   import ControlImage

##############################################
#Becario Digital con GUI                    ##
#Version 1.3                                ##
#La papeleria es lo mio, es mi razon de ser.##
# Cambios:                                  ##
# -botones en cada pestaña                  ##
# -partograma ampliado                      ##
# -fecha en algunos documentos              ##
##############################################

#Diccionarios
datospaciente = {'nombre': "",'edad': "",'FN': "",'CURP': "",'domicilio': "",'colonia': "",'cp': "",'municipio': "",'estado': "",'diagnosticos': "",'mpf': "",'medico': ""}
iden = []
cons_imprimir = {}
par_gram = {'gestas': "",'abortos': "",'partos': "",'cesareas': "", 'fum': "", 'sangre': "",'fpp': "",'amenorrea': 0}
fechahoy = datetime.now().replace(microsecond=0)

#Interfaz Grafica
class BecarioDigital(BaseWidget):

    def __init__(self):
        super(BecarioDigital,self).__init__('Becario Digital v1.3 - Edicion Corvid-19')
        formSettings.PYFORMS_STYLESHEET = './formatos/style.css'
        formSettings.PYFORMS_CONTROL_CODE_EDITOR_DEFAULT_FONT_SIZE = '12'
        
        #Definir campos de texto para el input
        self._nombre  = ControlText('Nombre')
        self._edad = ControlText('Edad')
        self._fn   = ControlText('Fecha de Nacimiento')
        self._CURP   = ControlText('CURP')
        self._domicilio   = ControlText('Domicilio')
        self._colonia   = ControlText('Colonia')
        self._cp   = ControlText('CP')
        self._municipio   = ControlText('Municipio')
        self._estado   = ControlText('Estado')
        self._diagnosticos   = ControlText('Diagnosticos\n(separados por comas o signo "+")')
        self._gestas   = ControlText('Gestas')
        self._partos   = ControlText('Partos')
        self._abortos   = ControlText('Abortos')
        self._cesareas   = ControlText('Cesareas')
        self._fum = ControlText('Fecha de Ultima Regla\n(dd/mm/aaaa)')
        self._sangre   = ControlText('Gpo y Rh')
        self._mpf   = ControlText('MPF')
        self._medico   = ControlText('Medico Especialista')
        self._button1  = ControlButton('Generar')
        self._button2  = ControlButton('Imprimir')

        #Definir checkboxs para imprimir consentimientos
        self._hojafrontal = ControlCheckBox('Hoja Frontal')
        self._partograma = ControlCheckBox('Partograma/Terminacion de Embarazo (3 Hojas)')
        self._ingresohospitalario = ControlCheckBox('Consentimiento Ingreso Hospitalario')
        self._atenciondelparto = ControlCheckBox('Consentimiento Atencion de parto')
        self._cesarea = ControlCheckBox('Consentimiento Cesarea (2 Hojas)')
        self._inducciondeparto = ControlCheckBox('Consentimiento Induccion de parto (2 Hojas)')
        self._anticoncepcion = ControlCheckBox('Consentimiento Anticoncepcion')
        self._interquirurgica = ControlCheckBox('Consentimiento Intervencion Quirurgica')
        self._LAPE = ControlCheckBox('Consentimiento LAPE (2 Hojas)')
        self._legradouterino = ControlCheckBox('Consentimiento Legrado Uterino (2 Hojas)')
        self._OTB = ControlCheckBox('Consentimiento OTB')
        
        #Definir botones generar/imprimir
        self._button1.value = self.__buttonAction1
        self._button2.value = self.__buttonAction2

        #Dar formato a los campos de texto
        self.formset = [ {
        'b:GenerarFormatos':['_nombre','||','_edad','||','_fn','=','_CURP','||','_domicilio','=','_colonia','||','_cp','||','_municipio','||','_estado','=','_diagnosticos','=',
        '_gestas','||','_partos','||','_abortos','||','_cesareas','=','_fum','||','_sangre','||','_mpf','||','_medico','=',('_button1')],
        'c:ImprimirFormatos':['_hojafrontal','||','_partograma','=','_ingresohospitalario','||','_atenciondelparto','=','_cesarea','||','_inducciondeparto','=','_anticoncepcion','||',
        '_interquirurgica','=','_LAPE','||','_legradouterino','=','_OTB','=',('_button2')],
        'a:Instrucciones':['Instrucciónes:','1-Seleccionar pestaña GenerarFormatos','2-Ingresa los datos necesarios','3-Dar click en boton Generar','4-Seleccionar pestaña ImprimirFormatos',
        '5-Dar click en la casilla de la papeleria a imprimir (Recomiendo una a la vez, se pueden seleccionar varias)','6-Dar click en boton imprimir','Repertir cuantas veces sea necesario',
        '=','Hecho por Diego Gonzalez']
        }]

    
    
    #Funcion para generar formatos apartir del input del usuario
    def __buttonAction1(self):
        datospaciente['nombre'] = self._nombre.value
        datospaciente['edad'] = self._edad.value
        datospaciente['FN'] = self._fn.value
        datospaciente['CURP'] = self._CURP.value
        datospaciente['domicilio'] = self._domicilio.value
        datospaciente['colonia'] = self._colonia.value
        datospaciente['cp'] = self._cp.value
        datospaciente['municipio'] = self._municipio.value
        datospaciente['estado'] = self._estado.value
        datospaciente['diagnosticos'] = self._diagnosticos.value
        datospaciente['mpf'] = self._mpf.value
        datospaciente['medico'] = self._medico.value
        par_gram['gestas'] = self._gestas.value
        par_gram['partos'] = self._partos.value
        par_gram['abortos'] = self._abortos.value
        par_gram['cesareas'] = self._cesareas.value
        fumc = self._fum.value
        par_gram['fum'] = fumc.replace(',','/').replace('-','/').replace('.','/')
        par_gram['sangre'] = self._sangre.value
        generar_formatos()

    #Funcion para imprimir formatos segun seleccion del usuario
    def __buttonAction2(self):
        cons_imprimir['hojafrontal'] = self._hojafrontal.value
        cons_imprimir['partograma'] = self._partograma.value
        cons_imprimir['ingresohospitalario'] = self._ingresohospitalario.value
        cons_imprimir['atenciondelparto'] = self._atenciondelparto.value
        cons_imprimir['cesarea'] = self._cesarea.value
        cons_imprimir['inducciondeparto'] = self._inducciondeparto.value
        cons_imprimir['anticoncepcion'] = self._anticoncepcion.value
        cons_imprimir['interquirurgica'] = self._interquirurgica.value
        cons_imprimir['LAPE'] = self._LAPE.value
        cons_imprimir['legradouterino'] = self._legradouterino.value
        cons_imprimir['OTB'] = self._OTB.value

        #Imprimir usando win32api
        for consentimiento in cons_imprimir:
            if cons_imprimir[consentimiento]:
                win32api.ShellExecute (0, "print", consentimiento + '.docx', None, ".", 0)
                time.sleep(5)


#lista de documentos a utilizar
consentimientos = os.listdir('./formatos/consentimientos/')


#Funcion principal
def generar_formatos():
    if par_gram['fum'] != "":
        fechaparto()
    con_info()
    partograma()
    hojafrontal()
    anticoncepcion()

#Funcion calcular fecha probable de parto y amenorrea usando el metodo Naegele
def fechaparto():
    FUM = datetime.strptime(par_gram['fum'],'%d/%m/%Y').date()
    duracion_gestacion_naegele = timedelta(days=280)
    sdg = date.today()
    amenorread = sdg - FUM
    par_gram['fpp'] = FUM + duracion_gestacion_naegele
    par_gram['amenorrea'] = int(amenorread.days) / 7

#Funcion para llenar consentimientos
def con_info():
    for consentimiento in consentimientos:
        document = Document('./formatos/consentimientos/' + consentimiento)
        tbl = document.tables[0]
        for rw in tbl.rows:
            iden.append(rw.cells[0].text)

        iden[1] = "El que suscribe:\nNombre Completo: {}               Fecha de Nacimiento: {} ".format(datospaciente['nombre'],datospaciente['FN'])
        iden[2] = "Edad: {}  Sexo: FEM    C.U.R.P:: {}      N° Expediente: SE    N° de S.P.: INSABI".format(datospaciente['edad'],datospaciente['CURP'])
        iden[4] = "Domicilio: {}             Colonia: {} ".format(datospaciente['domicilio'],datospaciente['colonia'])  
        iden[5] = "C.P: {}       Localidad:           Municipio: {}          Estado: {} ".format(datospaciente['cp'],datospaciente['municipio'],datospaciente['estado'])

        i=0
        for rw in tbl.rows:
            rw.cells[0].text = iden[i]
            i+= 1
            firmas = document.tables[1]
            firmas.rows[0].cells[0].text = "\t\t\t\tPaciente o Persona Legalmente Responsable\n\n\t\t\t\t{}".format(datospaciente['nombre'])
            
            if consentimiento == 'ingresohospitalario.docx':
                firmas.rows[2].cells[0].text = "\t\t\t\tMédico que realizara el procedimiento\n\n\t\t\t\t{}".format(datospaciente['medico'])
            
            document.save('./' + consentimiento)
            

#Para llenar el partograma, por ahora solo cambia la primer tabla
def partograma():
    document = Document('./formatos/partograma.docx')
    tbl = document.tables[0]
    tbl2 = document.tables[7]
    tbl3 = document.tables[5]
    tbl4 = document.tables[1]
    tbl5 = document.tables[4]
    tbl.rows[1].cells[0].text = "Nombre Completo: {}              Fecha de Nacimiento: {}  ".format(datospaciente['nombre'],datospaciente['FN'])
    tbl.rows[2].cells[0].text = "Edad: {}   Sexo: Fem  C.U.R.P.: {}  N° Exp:      N° S.P: INSABI".format(datospaciente['edad'],datospaciente['CURP'])
    tbl2.rows[1].cells[0].text = "Nombre Completo: {}              Fecha de Nacimiento: {}  ".format(datospaciente['nombre'],datospaciente['FN'])
    tbl3.rows[1].cells[0].text = "Nombre Completo: {}              Fecha de Nacimiento: {}  ".format(datospaciente['nombre'],datospaciente['FN'])
    tbl4.rows[2].cells[0].text = "Gesta: {}        Abortos: {}         Partos: {}         Cesareas: {}   ".format(par_gram['gestas'],par_gram['abortos'],par_gram['partos'],par_gram['cesareas'])
    tbl4.rows[3].cells[0].text = "FUM {}      FPP {}         Amenorrea {:.1f} sdg       MPF {}         Gpo y Rh {} ".format(par_gram['fum'],par_gram['fpp'],par_gram['amenorrea'],datospaciente['mpf'],par_gram['sangre'])
    tbl5.rows[2].cells[2].text = "      {}".format(datospaciente['medico'])
    document.save('./partograma.docx')

#llenar hoja frontal
def hojafrontal():
    document = Document('./formatos/hojafrontal.docx')
    tbl = document.tables[0]
    tbl2 = document.tables[1]
    tbl.rows[1].cells[0].text = "Nombre Completo: {}              Fecha de Nacimiento: {}  ".format(datospaciente['nombre'],datospaciente['FN'])
    tbl.rows[2].cells[0].text = "Edad: {}   Sexo: Fem  C.U.R.P.: {}  N° Exp:      N° S.P: INSABI".format(datospaciente['edad'],datospaciente['CURP'])
    tbl2.rows[1].cells[1].text = "{}".format(fechahoy)
    tbl2.rows[1].cells[2].text = "{}".format(datospaciente['diagnosticos'].replace(',','\n').replace('+','\n'))
    document.save('./hojafrontal.docx') 

#llenar hoja de anticoncepcion
def anticoncepcion():
    document = Document('./formatos/anticoncepcion.docx')
    paras = document.paragraphs
    hoy = date.today()
    paras[12].text = "Fecha:        {}            lugar: {}".format(hoy,datospaciente['estado'])
    paras[14].text = "La que suscribe: {}                y Seguro Popular  INSABI".format(datospaciente['nombre'])
    if datospaciente['mpf'] == "OTB" or datospaciente['mpf'] == "otb":
        paras[21].text = "Sin presión alguna, solicito y Autorizo al personal de salud de esta Unidad para que se me realice: Aplicación DE:    , método anticonceptivo temporal Realización de:  {}, método anticonceptivo definitivo".format(datospaciente['mpf'])
    else:
        paras[21].text = "Sin presión alguna, solicito y Autorizo al personal de salud de esta Unidad para que se me realice: Aplicación DE:  {}  , método anticonceptivo temporal Realización de:    , método anticonceptivo definitivo".format(datospaciente['mpf'])
    document.save('./anticoncepcion.docx') 

#Ejectuar aplicacion / tamaño de ventana
if __name__ == "__main__": pyforms.start_app( BecarioDigital, geometry=(800,700,800,700) )

